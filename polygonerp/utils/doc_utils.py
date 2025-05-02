import os

from docx import Document
from docx.enum.section import WD_ORIENTATION
from flask import current_app, send_file


def user_list_to_doc(users):
    doc = Document()
    doc.sections[0].orientation = WD_ORIENTATION.LANDSCAPE

    new_width, new_height = doc.sections[0].page_height, doc.sections[0].page_width
    doc.sections[0].page_width = new_width
    doc.sections[0].page_height = new_height

    doc.add_heading('User List', 0)
    table = doc.add_table(rows=1, cols=8, style='Table Grid')


    cells = table.rows[0].cells

    cells[0].text = "Name"
    cells[1].text = "Maiden name"
    cells[2].text = "Place of birth"
    cells[3].text = "Birth date"
    cells[4].text = "Mother's name"
    cells[5].text = "Address"
    cells[6].text = "Tax number"
    cells[7].text = "Taj number"

    for user in users:
        row_cells = table.add_row().cells
        row_cells[0].text = user.name
        row_cells[1].text = user.maiden_name
        row_cells[2].text = user.pob
        row_cells[3].text = user.dob
        row_cells[4].text = user.mothers_name
        row_cells[5].text = user.address
        row_cells[6].text = user.tax_num
        row_cells[7].text = user.taj_number

    filename = "employee_list.docx"
    list_dir = os.path.join(current_app.root_path, 'static/docs')
    full_path = os.path.join(list_dir, filename)

    if not os.path.exists(full_path):
        print("Not found")

    doc.save(os.path.join(current_app.root_path, 'static', 'docs', 'employee_list.docx'))
    return send_file(full_path, as_attachment=True)

def time_log_to_doc(logs, user_name):
    doc = Document()
    doc.sections[0].orientation = WD_ORIENTATION.LANDSCAPE
    new_width, new_height = doc.sections[0].page_height, doc.sections[0].page_width
    doc.sections[0].page_width = new_width
    doc.sections[0].page_height = new_height
    doc.add_heading(f'{user_name} Time Sheet {logs[0].log_date}', 0)
    table = doc.add_table(rows=1, cols=4, style='Table Grid')
    cells = table.rows[0].cells
    cells[0].text = "start_time"
    cells[1].text = "finish_time"
    cells[2].text = "total_time"
    cells[3].text = "log_type"

    for log in logs:
        row_cells = table.add_row().cells
        row_cells[0].text = str(log.start_time)
        row_cells[1].text = str(log.finish_time)
        row_cells[2].text = str(log.total_time)
        row_cells[3].text = str(log.log_type)

    filename = f'{user_name}_time_sheet.docx'
    list_dir = os.path.join(current_app.root_path, 'static/docs')
    full_path = os.path.join(list_dir, filename)


    doc.save(os.path.join(current_app.root_path, 'static', 'docs', f'{user_name}_time_sheet.docx'))
    return send_file(full_path, as_attachment=True)


def create_contract(user):
    doc = Document(os.path.join(current_app.root_path, 'static', 'docs', 'employee_contract_template.docx'))

    placeholders = {
        '{name}': user.name,
        '{maiden_name}': user.maiden_name,
        '{mothers_name}': user.mothers_name,
        '{pob}': user.pob,
        '{dob}': str(user.dob),
        '{address}': str(user.address),
        '{tax_num}': str(user.tax_num),
        '{taj_number}': str(user.taj_number),
        '{job_title}': user.job_title,
        '{base_pay}': str(user.base_pay)
    }

    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    doc.save(os.path.join(current_app.root_path, 'static', 'docs', f'{user.name}{user.id}_contract.docx'))


def find_user_contract(user):
    filename = f"{user.name}{user.id}_contract.docx"
    contracts_dir = os.path.join(current_app.root_path, 'static/docs')
    full_path = os.path.join(contracts_dir, filename)

    if not os.path.exists(full_path):
        raise FileNotFoundError("Cannot find contract file!")

    return send_file(full_path, as_attachment=True)